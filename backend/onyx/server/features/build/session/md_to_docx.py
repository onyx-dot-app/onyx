"""Convert Markdown to a DOCX document using mistune + python-docx.

Used by the build session "export as DOCX" feature. ``mistune`` parses the
Markdown into an AST and ``python-docx`` writes the ``.docx``; both are
pure-Python, so the conversion needs no external binary.

Supported constructs (covering what the LLM-generated documents emit):
headings, bold/italic/strikethrough/inline-code, bulleted/numbered/nested
lists (with loose-list continuation paragraphs and preserved ordered-list
start values), blockquotes, fenced code blocks, GFM tables, hyperlinks
(carrying inherited inline formatting), images (rendered as alt text), inline
``<br>`` line breaks, HTML entities, and horizontal rules. Other raw HTML is
dropped rather than shown as literal markup.
"""

from dataclasses import dataclass
from dataclasses import replace
from html import unescape
from io import BytesIO
from typing import Any
from typing import cast

import mistune
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.styles.style import ParagraphStyle
from docx.table import _Cell
from docx.text.paragraph import Paragraph

_MONOSPACE_FONT = "Courier New"
_CODE_FONT_SIZE = Pt(9)
_LINK_COLOR = "0563C1"
# python-docx ships built-in "List Bullet"/"List Number" styles plus numbered
# variants up to level 3 ("List Bullet 2", "List Bullet 3", ...). Deeper nesting
# reuses the level-3 style.
_MAX_LIST_LEVEL = 3

# Paragraph styles, mirroring how pandoc's default reference.docx names and
# spaces its prose so the output reads like the previous pandoc export. The
# spacing/indent/heading values below are reproduced from that reference (plain
# measurements, not the file itself, which stays out of the repo for licensing).
_STYLE_BODY = "Body Text"
_STYLE_FIRST_PARAGRAPH = "First Paragraph"
_STYLE_COMPACT = "Compact"
_STYLE_IMAGE_CAPTION = "Image Caption"
_STYLE_BLOCK_TEXT = "Block Text"

_BODY_SPACE = Pt(9)  # Body Text: 180 twips before/after
_COMPACT_SPACE = Pt(1.8)  # Compact (tight lists): 36 twips
_BLOCK_TEXT_SPACE = Pt(5)  # Block Text (blockquote): 100 twips
_BLOCK_TEXT_INDENT = Inches(1 / 3)  # Block Text left/right: 480 twips
_HEADING_COLOR = RGBColor(0x0F, 0x47, 0x61)
_HEADING_SIZES = {1: Pt(20), 2: Pt(16), 3: Pt(14), 4: Pt(12), 5: Pt(11), 6: Pt(11)}

# Enable the GFM table/strikethrough/url plugins so the AST covers the Markdown
# features that appear in these documents.
_markdown_parser = mistune.create_markdown(
    renderer=None,
    plugins=["table", "strikethrough", "url"],
)

Node = dict[str, Any]


@dataclass(frozen=True)
class _Fmt:
    """Inline formatting flags carried down through nested inline nodes."""

    bold: bool = False
    italic: bool = False
    strike: bool = False
    code: bool = False


def markdown_to_docx_bytes(md_text: str) -> bytes:
    """Render Markdown text to the bytes of a .docx file."""
    tokens = _markdown_parser(md_text)
    nodes: list[Node] = tokens if isinstance(tokens, list) else []

    document = Document()
    _apply_pandoc_styles(document)
    _render_blocks(document, nodes)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _apply_pandoc_styles(document: DocxDocument) -> None:
    """Add/configure the prose styles, approximating pandoc's reference.docx.

    python-docx's bare default template puts everything in ``Normal``; pandoc
    instead distributes prose across ``Body Text``/``First Paragraph``, tight
    lists into ``Compact``, blockquotes into ``Block Text``, and image captions
    into ``Image Caption``. Defining the same styles here lets the renderer
    assign them so the document reads like the pandoc export.
    """
    styles = document.styles
    existing = {style.name for style in styles}

    def ensure(name: str, base: str) -> ParagraphStyle:
        if name not in existing:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles[base]
            existing.add(name)
        return cast(ParagraphStyle, styles[name])

    body = cast(ParagraphStyle, styles[_STYLE_BODY])  # ships in the default template
    body.paragraph_format.space_before = _BODY_SPACE
    body.paragraph_format.space_after = _BODY_SPACE

    ensure(_STYLE_FIRST_PARAGRAPH, _STYLE_BODY)

    compact = ensure(_STYLE_COMPACT, _STYLE_BODY)
    compact.paragraph_format.space_before = _COMPACT_SPACE
    compact.paragraph_format.space_after = _COMPACT_SPACE

    block_text = ensure(_STYLE_BLOCK_TEXT, _STYLE_BODY)
    block_text.paragraph_format.space_before = _BLOCK_TEXT_SPACE
    block_text.paragraph_format.space_after = _BLOCK_TEXT_SPACE
    block_text.paragraph_format.left_indent = _BLOCK_TEXT_INDENT
    block_text.paragraph_format.right_indent = _BLOCK_TEXT_INDENT

    caption = ensure(_STYLE_IMAGE_CAPTION, "Caption")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for level, size in _HEADING_SIZES.items():
        heading = styles[f"Heading {level}"]
        heading.font.size = size
        heading.font.color.rgb = _HEADING_COLOR


# --------------------------------------------------------------------------- #
# Block-level rendering
# --------------------------------------------------------------------------- #
def _render_blocks(document: DocxDocument, nodes: list[Node]) -> None:
    # Like pandoc: the first prose paragraph after any non-paragraph block (a
    # heading, list, table, blockquote, code, or the document start) uses "First
    # Paragraph"; consecutive prose paragraphs use "Body Text".
    first_para_pending = True
    for node in nodes:
        node_type = node.get("type")
        if node_type in ("blank_line", "newline"):
            continue
        if node_type in ("paragraph", "block_text"):
            children = node.get("children", [])
            if _is_image_only(children):
                # An image renders as a caption; pandoc follows it with Body Text.
                _render_image_caption(document, children)
                first_para_pending = False
            else:
                style = _STYLE_FIRST_PARAGRAPH if first_para_pending else _STYLE_BODY
                paragraph = document.add_paragraph(style=style)
                _add_runs(paragraph, children, _Fmt())
                first_para_pending = False
            continue

        if node_type == "heading":
            level = min(int(node.get("attrs", {}).get("level", 1)), 6)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _add_runs(paragraph, node.get("children", []), _Fmt())
        elif node_type == "block_code":
            _render_code(document, node)
        elif node_type == "block_quote":
            _render_quote(document, node)
        elif node_type == "list":
            _render_list(document, node, level=0)
        elif node_type == "thematic_break":
            _render_thematic_break(document)
        elif node_type == "table":
            _render_table(document, node)
            # pandoc styles the paragraph after a table as Body Text, not First.
            first_para_pending = False
            continue
        elif "children" in node:
            # Unknown block wrapper: recurse so its content is not dropped.
            _render_blocks(document, node["children"])
        first_para_pending = True


def _render_code(document: DocxDocument, node: Node) -> None:
    raw = str(node.get("raw", "")).rstrip("\n")
    paragraph = document.add_paragraph()
    for index, line in enumerate(raw.split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = _MONOSPACE_FONT
        run.font.size = _CODE_FONT_SIZE


def _render_quote(document: DocxDocument, node: Node) -> None:
    for child in node.get("children", []):
        if child.get("type") == "paragraph":
            paragraph = document.add_paragraph(style=_STYLE_BLOCK_TEXT)
            _add_runs(paragraph, child.get("children", []), _Fmt())
        else:
            _render_blocks(document, [child])


def _is_image_only(children: list[Node]) -> bool:
    """True if the inline content is a single image (a standalone figure)."""
    meaningful = [
        child
        for child in children
        if child.get("type") not in ("softbreak", "linebreak")
        and not (child.get("type") == "text" and not str(child.get("raw", "")).strip())
    ]
    return len(meaningful) == 1 and meaningful[0].get("type") == "image"


def _render_image_caption(document: DocxDocument, children: list[Node]) -> None:
    """Render a standalone image as its alt text in the Image Caption style.

    Remote images are not fetched/embedded (pandoc does not either); the alt
    text is what reads in the document.
    """
    image = next(child for child in children if child.get("type") == "image")
    alt = _collect_text(image.get("children", []))
    paragraph = document.add_paragraph(style=_STYLE_IMAGE_CAPTION)
    paragraph.add_run(alt or "image")


def _render_list(document: DocxDocument, node: Node, level: int) -> None:
    attrs = node.get("attrs", {})
    ordered = bool(attrs.get("ordered", False))
    # mistune exposes tight/loose as a top-level key on the list node.
    tight = bool(node.get("tight", True))
    style_level = min(level + 1, _MAX_LIST_LEVEL)
    list_style = "List Number" if ordered else "List Bullet"
    if style_level > 1:
        list_style = f"{list_style} {style_level}"
    continue_style = (
        "List Continue" if style_level == 1 else f"List Continue {style_level}"
    )
    start = int(attrs.get("start", 1))

    # Tight lists match pandoc's "Compact" style. Compact carries no list marker,
    # so numbering is applied directly from the built-in list style's definition
    # (which also supplies the indentation); each list gets its own instance so
    # ordered lists restart correctly. Loose lists keep the built-in list style.
    if tight:
        num_id = _create_list_numbering(document, list_style, start)
        item_style = _STYLE_COMPACT if num_id is not None else list_style
    else:
        item_style = list_style
        num_id = (
            _create_list_numbering(document, list_style, start)
            if ordered and start != 1
            else None
        )

    for item in node.get("children", []):
        if item.get("type") != "list_item":
            continue
        has_rendered_marker = False
        for child in item.get("children", []):
            child_type = child.get("type")
            if child_type in ("blank_line", "newline"):
                continue
            if child_type in ("block_text", "paragraph"):
                paragraph_style = item_style if not has_rendered_marker else continue_style
                paragraph = document.add_paragraph(style=paragraph_style)
                if not has_rendered_marker and num_id is not None:
                    _apply_numbering(paragraph, num_id)
                _add_runs(paragraph, child.get("children", []), _Fmt())
                has_rendered_marker = True
            elif child_type == "list":
                _render_list(document, child, level + 1)
            else:
                _render_blocks(document, [child])


def _create_list_numbering(
    document: DocxDocument, list_style_name: str, start: int
) -> int | None:
    """Create a fresh numbering instance bound to a built-in list style.

    Returning a dedicated ``numId`` lets a paragraph in a non-list style (e.g.
    ``Compact``) still render list markers + indentation, and gives each list an
    independent counter so ordered lists restart. ``start`` adds a
    ``startOverride`` when the list does not begin at 1. Returns None if the
    style has no numbering definition (caller falls back to the list style).
    """
    style = document.styles[list_style_name]
    numbering = document.part.numbering_part.element
    abstract_num_id = _abstract_num_id_for_style(numbering, style.style_id)
    if abstract_num_id is None:
        return None

    num_ids = [
        int(num.get(qn("w:numId")))
        for num in numbering.findall(qn("w:num"))
        if num.get(qn("w:numId")) is not None
    ]
    next_num_id = max(num_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))

    abstract_num_id_el = OxmlElement("w:abstractNumId")
    abstract_num_id_el.set(qn("w:val"), abstract_num_id)
    num.append(abstract_num_id_el)

    if start != 1:
        lvl_override = OxmlElement("w:lvlOverride")
        lvl_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), str(start))
        lvl_override.append(start_override)
        num.append(lvl_override)

    numbering.append(num)
    return next_num_id


def _abstract_num_id_for_style(numbering: Any, style_id: str) -> str | None:
    for abstract_num in numbering.findall(qn("w:abstractNum")):
        for lvl in abstract_num.findall(qn("w:lvl")):
            p_style = lvl.find(qn("w:pStyle"))
            if p_style is not None and p_style.get(qn("w:val")) == style_id:
                return abstract_num.get(qn("w:abstractNumId"))
    return None


def _apply_numbering(paragraph: Paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = num_pr.get_or_add_ilvl()
    ilvl.val = 0
    num_id_el = num_pr.get_or_add_numId()
    num_id_el.val = num_id


def _render_thematic_break(document: DocxDocument) -> None:
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    borders.append(bottom)
    p_pr.append(borders)


def _render_table(document: DocxDocument, node: Node) -> None:
    header_cells: list[Node] = []
    body_rows: list[list[Node]] = []
    for section in node.get("children", []):
        section_type = section.get("type")
        if section_type == "table_head":
            header_cells = section.get("children", [])
        elif section_type == "table_body":
            for row in section.get("children", []):
                body_rows.append(row.get("children", []))

    num_cols = len(header_cells) or (len(body_rows[0]) if body_rows else 0)
    if num_cols == 0:
        return

    table = document.add_table(rows=0, cols=num_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    if header_cells:
        cells = table.add_row().cells
        for index, cell_node in enumerate(header_cells[:num_cols]):
            _fill_cell(cells[index], cell_node, bold=True)
    for row in body_rows:
        cells = table.add_row().cells
        for index, cell_node in enumerate(row[:num_cols]):
            _fill_cell(cells[index], cell_node, bold=False)


def _fill_cell(cell: _Cell, cell_node: Node, bold: bool) -> None:
    paragraph = cell.paragraphs[0]
    _add_runs(paragraph, cell_node.get("children", []), _Fmt(bold=bold))


# --------------------------------------------------------------------------- #
# Inline rendering
# --------------------------------------------------------------------------- #
def _add_runs(paragraph: Paragraph, nodes: list[Node], fmt: _Fmt) -> None:
    for node in nodes:
        node_type = node.get("type")
        if node_type == "text":
            # Decode HTML entities (e.g. ``&amp;``, ``&copy;``) the way a Markdown
            # renderer would; code spans below are intentionally left literal.
            _styled_run(paragraph, unescape(str(node.get("raw", ""))), fmt)
        elif node_type == "strong":
            _add_runs(paragraph, node.get("children", []), replace(fmt, bold=True))
        elif node_type == "emphasis":
            _add_runs(paragraph, node.get("children", []), replace(fmt, italic=True))
        elif node_type == "strikethrough":
            _add_runs(paragraph, node.get("children", []), replace(fmt, strike=True))
        elif node_type == "codespan":
            _styled_run(paragraph, str(node.get("raw", "")), replace(fmt, code=True))
        elif node_type == "link":
            _add_hyperlink(paragraph, node, fmt)
        elif node_type == "image":
            alt = _collect_text(node.get("children", []))
            _styled_run(paragraph, f"[image: {alt}]" if alt else "[image]", fmt)
        elif node_type == "softbreak":
            _styled_run(paragraph, " ", fmt)
        elif node_type == "linebreak":
            paragraph.add_run().add_break()
        elif node_type == "inline_html":
            _add_inline_html(paragraph, str(node.get("raw", "")))
        elif "children" in node:
            _add_runs(paragraph, node["children"], fmt)
        elif "raw" in node:
            _styled_run(paragraph, unescape(str(node["raw"])), fmt)


def _styled_run(paragraph: Paragraph, text: str, fmt: _Fmt) -> None:
    run = paragraph.add_run(text)
    run.bold = fmt.bold
    run.italic = fmt.italic
    if fmt.strike:
        run.font.strike = True
    if fmt.code:
        run.font.name = _MONOSPACE_FONT


def _add_hyperlink(paragraph: Paragraph, node: Node, fmt: _Fmt) -> None:
    url = str(node.get("attrs", {}).get("url", ""))
    text = _collect_text(node.get("children", [])) or url
    if not url:
        _styled_run(paragraph, text, fmt)
        return

    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    # Children must follow the CT_RPr schema order (rFonts, b, i, strike, color,
    # u) so the inherited formatting carried in ``fmt`` survives on the link run.
    if fmt.code:
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), _MONOSPACE_FONT)
        fonts.set(qn("w:hAnsi"), _MONOSPACE_FONT)
        run_props.append(fonts)
    if fmt.bold:
        run_props.append(OxmlElement("w:b"))
    if fmt.italic:
        run_props.append(OxmlElement("w:i"))
    if fmt.strike:
        run_props.append(OxmlElement("w:strike"))
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _LINK_COLOR)
    run_props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)
    run.append(run_props)
    text_el = OxmlElement("w:t")
    text_el.text = text
    # Preserve leading/trailing whitespace; OOXML strips it without this hint.
    if text != text.strip():
        text_el.set(qn("xml:space"), "preserve")
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline_html(paragraph: Paragraph, raw: str) -> None:
    normalized = raw.strip().lower()
    if normalized in ("<br>", "<br/>", "<br />"):
        paragraph.add_run().add_break()


def _collect_text(nodes: list[Node]) -> str:
    """Flatten inline nodes to plain text (for link labels and image alt text)."""
    parts: list[str] = []
    for node in nodes:
        if node.get("type") == "text":
            parts.append(unescape(str(node.get("raw", ""))))
        elif node.get("children"):
            parts.append(_collect_text(node["children"]))
        elif "raw" in node:
            parts.append(unescape(str(node["raw"])))
    return "".join(parts)
