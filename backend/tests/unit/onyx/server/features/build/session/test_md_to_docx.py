"""Unit tests for the Markdown -> DOCX converter."""

import re
import zipfile
from io import BytesIO

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn

from onyx.server.features.build.session.md_to_docx import markdown_to_docx_bytes


def _render(md_text: str) -> DocxDocument:
    data = markdown_to_docx_bytes(md_text)
    # A .docx is a zip archive, which always starts with the "PK" magic bytes.
    assert data[:2] == b"PK"
    return Document(BytesIO(data))


def _paragraph_styles(doc: DocxDocument) -> list[str]:
    return [
        p.style.name
        for p in doc.paragraphs
        if p.text.strip() and p.style is not None and p.style.name is not None
    ]


def test_empty_markdown_produces_valid_docx() -> None:
    doc = _render("")
    assert doc.tables == []


def test_headings_map_to_heading_styles() -> None:
    doc = _render("# H1\n\n## H2\n\n### H3\n")
    styles = _paragraph_styles(doc)
    assert "Heading 1" in styles
    assert "Heading 2" in styles
    assert "Heading 3" in styles
    texts = [p.text for p in doc.paragraphs]
    assert "H1" in texts and "H2" in texts and "H3" in texts


def test_inline_formatting_sets_run_properties() -> None:
    doc = _render("Plain **bold** and *italic* and ~~strike~~ and `code`.")
    runs = [r for p in doc.paragraphs for r in p.runs]
    assert any(r.text == "bold" and r.bold for r in runs)
    assert any(r.text == "italic" and r.italic for r in runs)
    assert any(r.text == "strike" and r.font.strike for r in runs)
    assert any(r.text == "code" and r.font.name == "Courier New" for r in runs)


def test_tight_lists_use_compact_style_with_numbering() -> None:
    # Tight lists match pandoc's "Compact" style; markers come from direct
    # numbering (numPr) rather than a list paragraph style.
    md = "- a\n- b\n  - nested\n\n1. one\n2. two\n"
    doc = _render(md)
    assert set(_paragraph_styles(doc)) == {"Compact"}
    list_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    assert all(p._p.find(".//" + qn("w:numPr")) is not None for p in list_paragraphs)


def test_loose_lists_keep_list_paragraph_styles() -> None:
    # A blank line between items makes the list loose; pandoc keeps the built-in
    # numbered/bulleted list paragraph styles there.
    doc = _render("1. one\n\n2. two\n")
    assert "List Number" in _paragraph_styles(doc)


def test_loose_list_continuation_paragraphs_are_not_numbered() -> None:
    doc = _render("1. one\n\n   continuation\n2. two\n")
    non_empty_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    assert [p.text for p in non_empty_paragraphs] == ["one", "continuation", "two"]
    assert [
        p.style.name
        for p in non_empty_paragraphs
        if p.style is not None and p.style.name is not None
    ] == [
        "List Number",
        "List Continue",
        "List Number",
    ]


def test_ordered_list_start_value_is_preserved() -> None:
    doc = _render("3. three\n4. four\n")
    assert [p.text for p in doc.paragraphs if p.text.strip()] == ["three", "four"]
    numbering_xml = doc.part.numbering_part.element.xml
    assert '<w:startOverride w:val="3"/>' in numbering_xml


def test_block_quote_uses_block_text_style() -> None:
    # Matches pandoc, which puts blockquotes in the indented "Block Text" style.
    doc = _render("> quoted line\n")
    assert "Block Text" in _paragraph_styles(doc)


def test_prose_uses_body_text_and_first_paragraph_styles() -> None:
    # The first paragraph after a heading is "First Paragraph"; the next is
    # "Body Text" (pandoc's convention), and neither falls back to "Normal".
    doc = _render("# Title\n\nFirst para.\n\nSecond para.\n")
    styles = _paragraph_styles(doc)
    assert styles == ["Heading 1", "First Paragraph", "Body Text"]


def test_standalone_image_becomes_image_caption() -> None:
    doc = _render("![Asimov portrait](https://example.com/a.png)\n")
    captions = [
        p.text for p in doc.paragraphs if p.style and p.style.name == "Image Caption"
    ]
    # Alt text is shown (not embedded), with no "[image: ...]" wrapper.
    assert captions == ["Asimov portrait"]
    assert "[image" not in doc.paragraphs[0].text


def test_fenced_code_block_is_monospace() -> None:
    doc = _render("```python\ndef f():\n    return 1\n```\n")
    monospace_runs = [
        r.text for p in doc.paragraphs for r in p.runs if r.font.name == "Courier New"
    ]
    assert any("def f():" in t for t in monospace_runs)


def test_table_is_rendered_with_header_and_rows() -> None:
    md = "| Name | Value |\n|------|-------|\n| a | 1 |\n| b | 2 |\n"
    doc = _render(md)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert (len(table.rows), len(table.columns)) == (3, 2)
    assert [c.text for c in table.rows[0].cells] == ["Name", "Value"]
    assert [c.text for c in table.rows[1].cells] == ["a", "1"]


def test_link_becomes_real_hyperlink() -> None:
    doc = _render("See [Onyx](https://onyx.app) for details.")
    xml = doc.element.xml
    assert "w:hyperlink" in xml
    # The visible link text is preserved.
    assert "Onyx" in xml


def test_link_without_url_falls_back_to_text() -> None:
    # An autolink-free bare reference still renders its text without crashing.
    doc = _render("[empty]()")
    assert any("empty" in p.text for p in doc.paragraphs)


def test_hyperlink_inherits_surrounding_formatting() -> None:
    # A link nested inside bold should keep the bold run property on the link.
    doc = _render("**[click here](https://onyx.app)**")
    xml = doc.element.xml
    assert "w:hyperlink" in xml
    hyperlink_run_props = xml.split("w:hyperlink")[1]
    assert "<w:b" in hyperlink_run_props


def test_hyperlink_preserves_whitespace_in_text() -> None:
    # Trailing space inside the link text must not be stripped by OOXML.
    doc = _render("a [ spaced ](https://onyx.app) b")
    assert 'xml:space="preserve"' in doc.element.xml


def test_inline_html_br_becomes_docx_line_break() -> None:
    doc = _render("before<br>after")
    assert [p.text for p in doc.paragraphs] == ["before\nafter"]
    assert "<w:br/>" in doc.element.xml


def test_unsupported_inline_html_tags_are_not_visible_text() -> None:
    doc = _render("<u>under</u> normal")
    assert [p.text for p in doc.paragraphs] == ["under normal"]
    assert "<u>" not in doc.paragraphs[0].text
    assert "</u>" not in doc.paragraphs[0].text


def test_html_entities_are_decoded_in_text() -> None:
    doc = _render("A &amp; B &lt; C &copy; 2026 &mdash; D")
    assert [p.text for p in doc.paragraphs] == ["A & B < C © 2026 — D"]


def test_html_entities_are_decoded_in_link_text() -> None:
    doc = _render("[R &amp; D](https://onyx.app)")
    xml = doc.element.xml
    # "&" is stored decoded (serialized as the single XML escape &amp;), not left
    # as the literal entity text (which would double-escape to &amp;amp;).
    assert "R &amp; D" in xml
    assert "&amp;amp;" not in xml


def test_html_entities_in_code_span_are_left_literal() -> None:
    # CommonMark does not decode entities inside code spans.
    doc = _render("`a &amp; b`")
    runs = [r.text for p in doc.paragraphs for r in p.runs]
    assert "a &amp; b" in runs


def _footnote_ids(data: bytes) -> list[str]:
    archive = zipfile.ZipFile(BytesIO(data))
    footnotes_xml = archive.read("word/footnotes.xml").decode()
    return re.findall(r'<w:footnote [^>]*w:id="(-?\d+)"', footnotes_xml)


def test_footnotes_become_real_word_footnotes() -> None:
    data = markdown_to_docx_bytes("A citation.[^1]\n\n[^1]: The note text.\n")
    archive = zipfile.ZipFile(BytesIO(data))
    # The body references a real footnote rather than literal "[^1]" text.
    assert "<w:footnoteReference" in archive.read("word/document.xml").decode()
    body_text = "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)
    assert "[^1]" not in body_text
    # The note content lives in the footnotes part (ids -1/0 are the separators).
    footnotes_xml = archive.read("word/footnotes.xml").decode()
    assert "The note text." in footnotes_xml
    assert _footnote_ids(data) == ["-1", "0", "1"]


def test_repeated_footnote_reference_emits_one_note_each() -> None:
    # Word footnotes are 1:1 with their reference, so a note cited twice produces
    # two footnotes (matching pandoc), not one shared id.
    data = markdown_to_docx_bytes("First.[^1] Second.[^1]\n\n[^1]: Shared note.\n")
    assert _footnote_ids(data) == ["-1", "0", "1", "2"]


def test_document_without_footnotes_has_no_footnotes_part() -> None:
    data = markdown_to_docx_bytes("Just a paragraph.\n")
    assert "word/footnotes.xml" not in zipfile.ZipFile(BytesIO(data)).namelist()


def test_default_font_matches_pandoc() -> None:
    # pandoc's reference uses Aptos 12pt body / Aptos Display headings, not
    # python-docx's Cambria 11pt default.
    doc = _render("# Heading\n\nBody.\n")
    normal = doc.styles["Normal"]
    assert normal.font.name == "Aptos"
    assert normal.font.size is not None and normal.font.size.pt == 12
    assert doc.styles["Heading 1"].font.name == "Aptos Display"


def test_page_margins_match_pandoc() -> None:
    # pandoc renders 1" margins (Word default); python-docx's template uses 1.25"
    # left/right, so set 1" all round to match.
    section = _render("Body.\n").sections[0]
    margins = (
        section.left_margin,
        section.right_margin,
        section.top_margin,
        section.bottom_margin,
    )
    assert all(margin is not None and margin.inches == 1.0 for margin in margins)


def test_paragraph_style_sequence_matches_pandoc_rules() -> None:
    """Lock in pandoc's paragraph-style assignment across mixed content.

    The expected styles below were verified against pandoc's own DOCX output:
    the first prose paragraph after any non-paragraph block uses "First
    Paragraph" -- except after a table, where it is "Body Text".
    """
    md = (
        "# Heading\n\n"
        "First body paragraph.\n\n"
        "Second body paragraph.\n\n"
        "- bullet one\n"
        "- bullet two\n\n"
        "Paragraph after list.\n\n"
        "> A blockquote.\n\n"
        "Paragraph after quote.\n\n"
        "| A | B |\n|---|---|\n| 1 | 2 |\n\n"
        "Paragraph after table.\n"
    )
    doc = _render(md)
    assert _paragraph_styles(doc) == [
        "Heading 1",
        "First Paragraph",  # first prose after a heading
        "Body Text",  # consecutive prose
        "Compact",  # tight list item
        "Compact",
        "First Paragraph",  # first prose after a list
        "Block Text",  # blockquote
        "First Paragraph",  # first prose after a blockquote
        "Body Text",  # prose after a table (the exception)
    ]
