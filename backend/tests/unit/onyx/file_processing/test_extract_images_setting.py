"""Regression tests for the "Extract & Caption Images" workspace setting.

`extract_text_and_images` is the extractor used by the file connector, which
backs both connector indexing *and* chat/project file uploads. Every
image-bearing branch in it must honour the workspace setting; previously only
the PDF branch did, so a DOCX/PPTX uploaded in chat had its embedded images
extracted and stored (and later captioned) even with the toggle off.

The docx/pptx fixtures are built in-memory so these tests need no new binary
fixtures: the docx is written with python-docx (already a hard dependency, and
markitdown must be able to parse it for the text branch), and the pptx is a
minimal zip since pptx image extraction just reads `ppt/media/*` out of it.
"""

import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any

import pytest

from onyx.file_processing import extract_file_text as extract_file_text_module
from onyx.file_processing.extract_file_text import extract_text_and_images

FIXTURES = Path(__file__).parent / "fixtures"


def _png_bytes(size: int = 32) -> bytes:
    """A small but real PNG. Generated rather than hard-coded so python-docx's
    image header parsing gets something genuinely valid."""
    from PIL import Image

    buf = BytesIO()
    Image.new("RGB", (size, size), (200, 30, 30)).save(buf, format="PNG")
    return buf.getvalue()


def _load_pdf() -> BytesIO:
    return BytesIO((FIXTURES / "with_image.pdf").read_bytes())


def _make_docx_with_image() -> BytesIO:
    """A real docx containing one embedded image, built with python-docx.

    Hand-rolling the zip is not enough here: markitdown parses the document for
    the text half of the same call, so the package has to be well-formed.
    """
    import docx
    from docx.shared import Inches

    document = docx.Document()
    document.add_paragraph("Hello docx")
    document.add_picture(BytesIO(_png_bytes()), width=Inches(1))

    buf = BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf


def _make_pptx_with_image() -> BytesIO:
    """Minimal pptx containing one embedded image under ppt/media/."""
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org'
            '/package/2006/content-types"><Default Extension="xml" ContentType='
            '"application/xml"/><Default Extension="png" ContentType="image/png"/>'
            "</Types>",
        )
        z.writestr(
            "_rels/.rels",
            '<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats'
            '.org/package/2006/relationships"/>',
        )
        z.writestr("ppt/media/image1.png", _png_bytes())
    buf.seek(0)
    return buf


@pytest.fixture(autouse=True)
def _no_unstructured(monkeypatch: pytest.MonkeyPatch) -> None:
    """Keep the Unstructured short-circuit out of the way."""
    monkeypatch.setattr(
        extract_file_text_module, "get_unstructured_api_key", lambda: None
    )


def _set_setting(monkeypatch: pytest.MonkeyPatch, enabled: bool) -> None:
    monkeypatch.setattr(
        extract_file_text_module,
        "get_image_extraction_and_analysis_enabled",
        lambda: enabled,
    )


# ── setting disabled: no images from any format ──────────────────────────


class TestImageExtractionDisabled:
    """With the toggle off, nothing image-bearing may come back — this is the
    path a chat PDF/DOCX/PPTX upload takes via LocalFileConnector."""

    def test_pdf_yields_no_images(self, monkeypatch: pytest.MonkeyPatch) -> None:
        _set_setting(monkeypatch, False)
        result = extract_text_and_images(_load_pdf(), "with_image.pdf")
        assert result.embedded_images == []

    def test_docx_yields_no_images(self, monkeypatch: pytest.MonkeyPatch) -> None:
        _set_setting(monkeypatch, False)
        result = extract_text_and_images(_make_docx_with_image(), "sample.docx")
        assert result.embedded_images == []

    def test_pptx_yields_no_images(self, monkeypatch: pytest.MonkeyPatch) -> None:
        _set_setting(monkeypatch, False)
        result = extract_text_and_images(_make_pptx_with_image(), "sample.pptx")
        assert result.embedded_images == []

    @pytest.mark.parametrize(
        ("file_factory", "file_name"),
        [
            (_load_pdf, "with_image.pdf"),
            (_make_docx_with_image, "sample.docx"),
            (_make_pptx_with_image, "sample.pptx"),
        ],
    )
    def test_image_callback_is_never_invoked(
        self,
        monkeypatch: pytest.MonkeyPatch,
        file_factory: Any,
        file_name: str,
    ) -> None:
        """The streaming path must be suppressed too. Connectors pass a callback
        that writes straight to the FileStore, so an un-gated callback would
        persist images even though `embedded_images` comes back empty."""
        _set_setting(monkeypatch, False)
        collected: list[tuple[bytes, str]] = []

        result = extract_text_and_images(
            file_factory(),
            file_name,
            image_callback=lambda data, name: collected.append((data, name)),
        )

        assert collected == []
        assert result.embedded_images == []

    def test_text_extraction_is_unaffected(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Disabling image extraction must not cost us the document text."""
        _set_setting(monkeypatch, False)
        result = extract_text_and_images(_load_pdf(), "with_image.pdf")
        assert isinstance(result.text_content, str)


# ── setting enabled: images still flow ───────────────────────────────────


class TestImageExtractionEnabled:
    def test_pdf_yields_images(self, monkeypatch: pytest.MonkeyPatch) -> None:
        _set_setting(monkeypatch, True)
        result = extract_text_and_images(_load_pdf(), "with_image.pdf")
        assert len(result.embedded_images) >= 1

    def test_docx_yields_images(self, monkeypatch: pytest.MonkeyPatch) -> None:
        _set_setting(monkeypatch, True)
        result = extract_text_and_images(_make_docx_with_image(), "sample.docx")
        assert len(result.embedded_images) >= 1

    def test_pptx_yields_images(self, monkeypatch: pytest.MonkeyPatch) -> None:
        _set_setting(monkeypatch, True)
        result = extract_text_and_images(_make_pptx_with_image(), "sample.pptx")
        assert len(result.embedded_images) >= 1

    def test_image_callback_still_streams(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        _set_setting(monkeypatch, True)
        collected: list[tuple[bytes, str]] = []

        result = extract_text_and_images(
            _load_pdf(),
            "with_image.pdf",
            image_callback=lambda data, name: collected.append((data, name)),
        )

        assert len(collected) >= 1
        # Streaming mode returns an empty list by design.
        assert result.embedded_images == []
